import json
import logging as log
import os
import pandas as pd
import shutil
import traceback
import email_utility
from os import listdir
from send_email import send_error_email
from docx import Document
from docx2pdf import convert
from docx.shared import Pt


def write_to_word(src, dest, params):
    """

    :param src:
    :param dest:
    :param params:
    :return:
    """
    for name in params.keys():
        doc_template_path = src
        pdf_path = dest
        user_doc_file = doc_template_path.split('\\')[-1]
        renamed_doc_file = '{}{}{}{}'.format(user_doc_file.split('_')[0].split('.')[0], '_',
                                             name.replace('/', '').replace(':', ''), '.docx')
        log.info(('Copying doc template from {} to {}'.format(doc_template_path, pdf_path)))
        shutil.copy(doc_template_path, pdf_path)
        os.chdir(pdf_path)
        log.info('Renaming the doc template by appending the username to it')
        renamed_doc_dest_path = os.path.join(os.getcwd(), renamed_doc_file)
        os.rename(os.path.join(pdf_path, user_doc_file), renamed_doc_dest_path)
        new_value = ""
        log.info('Reading the Document {} at path {} '.format(renamed_doc_file, renamed_doc_dest_path))
        doc = Document(renamed_doc_dest_path)
        for p in doc.paragraphs:
            if p.text.find("Account Name:") >= 0:
                old_value = p.text.split(':')[-1]
                p.text = p.text.replace(old_value, new_value)
                run = p.add_run(name)
                run.bold = True
                run.underline = True
                font = p.style.font
                font.size = Pt(12)
        log.info('Saving the content to file {}'.format(renamed_doc_file))
        doc.save(renamed_doc_dest_path)


def replace_spl_characters(str_):
    """

    :param str_:
    :return:
    """
    return str_.replace('/', '').replace(':', '')


def convert_word_to_pdf(src, params):
    """

    :param src:
    :param params:
    :return:
    """
    pdf_path = src
    doc_names = [replace_spl_characters(name) for name in params.keys()]
    for file_name in listdir(pdf_path):
        if file_name.endswith('.docx'):
            if file_name.split('.')[0].split('_')[-1] in doc_names:
                log.info('Converting word docx to pdf file')
                pdf_file_name = file_name.split('.')[0] + '.pdf'
                convert(pdf_path+file_name, pdf_path+pdf_file_name)
                os.remove(file_name)


def read_json(filepath):
    """

    :param filepath:
    :return:
    """
    with open(filepath, 'r') as jp:
        return json.load(jp)


def generate_html(header, html_body, signature):
    """

    :param header:
    :param html_body:
    :param signature:
    :return:
    """
    html = """
    <html>
    <body>
    <p>{0}</p>
    <p>{1}</p>
    <p>{2}</p>
    </body>
    </html>
    """.format(header, html_body, "<br>".join(signature.split(',')))
    return html


def get_complete_path(dir_, folder=None, file_name=None):
    """

    :param dir_:
    :param folder:
    :param file_name:
    :return:
    """
    if folder and file_name:
        return os.path.join(dir_, folder, file_name)
    elif file_name:
        return os.path.join(dir_, file_name)
    else:
        return os.path.join(dir_, folder)


if __name__ == "__main__":
    try:
        log.basicConfig(format='%(asctime)s - %(message)s', level=log.INFO)
        log.info('The program is executed by user {}'.format(os.getlogin()))
        abs_path = os.path.dirname(os.path.realpath(__file__))
        log.info('Fetching the doc_path, pdf_path and users info from the config file')
        os.chdir(abs_path)
        config_filename = get_complete_path(dir_=abs_path, folder="input", file_name="config.json")
        email_config_filename = get_complete_path(dir_=abs_path, folder="input"
                                                  , file_name="email_config.json")
        config = read_json(config_filename)
        email_config = read_json(email_config_filename)
        doc_path, pdf_path, csv_path = get_complete_path(dir_=abs_path, folder=config["doc_path"])\
            , get_complete_path(dir_=abs_path, folder=config["pdf_path"])\
            , get_complete_path(dir_=abs_path, folder=config["csv_path"])
        support_emails, from_email, email_subject, email_header = config["support_emails"]\
            , email_config["from_email"]\
            , email_config["email_subject"]\
            , email_config["email_body"]["header"]
        email_body, email_signature = email_config["email_body"]["body_content"]\
            , email_config["email_body"]["signature"]
        body = generate_html(header=email_header
                             , html_body=email_body
                             , signature=email_signature)
        csv_data = pd.read_csv(csv_path)
        dict_csv = {}
        for data in csv_data.values:
            user_name = data[0]
            email_id = data[1]
            dict_csv[user_name] = email_id
        write_to_word(src=doc_path, dest=pdf_path, params=dict_csv)
        convert_word_to_pdf(src=pdf_path, params=dict_csv)
        os.chdir(pdf_path)
        # email_utility.create_email(send_from=from_email,
        #                            subject=email_subject,
        #                            body=body,
        #                            params=dict_csv)
        log.info('doc_path={} pdf_path={}'.format(doc_path, pdf_path))
        log.info('End of the process to convert docx file to pdf file')
    except Exception as ex:
        print("Exception occurred: {}".format(ex))
        send_error_email(traceback=traceback.format_exc(),
                         subject="[WordToPDF] Runtime Error", receiver=support_emails)
